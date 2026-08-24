#!/usr/bin/env python3
"""Generate a professionally designed Word CV from a structured JSON file.

Usage:
  python scripts/generate_designed_cv.py input.json output.docx

The JSON schema is intentionally simple so Manus can create it after extracting a PDF,
LinkedIn export, biography, portfolio page, or raw notes.
"""

import json
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

NAVY = "0B1F33"
BLUE = "1F6FEB"
WHITE = "FFFFFF"
DARK = "111827"
GREY = "5B6770"
LIGHT_BLUE = "EAF3FF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ["top", "left", "bottom", "right"]:
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_empty_initial_paragraph(cell):
    if cell.paragraphs and not cell.paragraphs[0].text:
        p = cell.paragraphs[0]._element
        p.getparent().remove(p)


def run(paragraph, text, bold=False, size=9, color=DARK, italic=False, caps=False):
    r = paragraph.add_run(text.upper() if caps else text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Aptos"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    return r


def heading(cell, text, color=BLUE):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run(p, text, bold=True, size=10.5, color=color, caps=True)


def side_heading(cell, text):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run(p, text, bold=True, size=9.5, color=WHITE, caps=True)


def para(cell, text, size=9, color=DARK, bold=False, after=3):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    run(p, text, bold=bold, size=size, color=color)
    return p


def bullet(cell, text, size=8.5, color=DARK):
    p = cell.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(2)
    run(p, "• ", size=size, color=color)
    run(p, text, size=size, color=color)


def add_role(cell, role):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    run(p, role.get("title", ""), bold=True, size=10.2)
    org = role.get("organisation") or role.get("organization") or ""
    period = role.get("period", "")
    location = role.get("location", "")
    para(cell, org, size=8.8, color=BLUE, bold=True, after=1)
    meta = " · ".join([x for x in [period, location] if x])
    if meta:
        para(cell, meta, size=8.2, color=GREY, after=3)
    if role.get("description"):
        para(cell, role["description"], size=8.8, after=3)
    for item in role.get("bullets", []):
        bullet(cell, item)


def build_doc(data, output_path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    styles["Normal"].font.size = Pt(9)

    header = doc.add_table(rows=1, cols=1)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    hc = header.cell(0, 0)
    set_cell_shading(hc, NAVY)
    set_cell_margins(hc, top=260, bottom=230, start=260, end=260)
    set_cell_border(hc)
    remove_empty_initial_paragraph(hc)

    p = hc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(3)
    run(p, data.get("name", ""), bold=True, size=22, color=WHITE)
    p = hc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(4)
    run(p, data.get("headline", ""), size=10.5, color="CFE8FF")
    if data.get("tagline"):
        p = hc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, data["tagline"], size=9, color=WHITE)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.25)
    table.columns[1].width = Inches(4.65)
    left, right = table.cell(0, 0), table.cell(0, 1)
    left.width, right.width = Inches(2.25), Inches(4.65)
    set_cell_shading(left, NAVY)
    for c in [left, right]:
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_border(c)
        remove_empty_initial_paragraph(c)
    set_cell_margins(left, top=180, bottom=180, start=170, end=170)
    set_cell_margins(right, top=140, bottom=140, start=220, end=120)

    for section_name in ["contact", "core_profile", "skills", "languages", "awards", "certifications"]:
        items = data.get(section_name, [])
        if not items:
            continue
        label = section_name.replace("_", " ").title()
        side_heading(left, label)
        for item in items:
            para(left, str(item), size=8.3, color=WHITE, after=2)

    if data.get("profile"):
        heading(right, "Profil")
        for block in data["profile"] if isinstance(data["profile"], list) else [data["profile"]]:
            para(right, block, size=9.2, after=6)

    if data.get("strategic_focus"):
        heading(right, "Strategischer Fokus")
        for item in data["strategic_focus"]:
            bullet(right, item)

    if data.get("experience"):
        heading(right, "Berufserfahrung")
        for role in data["experience"]:
            add_role(right, role)

    if data.get("education"):
        heading(right, "Ausbildung")
        for edu in data["education"]:
            p = right.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(1)
            run(p, edu.get("degree", ""), bold=True, size=9.3)
            meta = " · ".join([x for x in [edu.get("institution", ""), edu.get("period", "")] if x])
            if meta:
                para(right, meta, size=8.4, color=GREY, after=2)

    if data.get("publications"):
        heading(right, "Publikationen und Themen")
        para(right, data["publications"], size=9.0, after=2)

    doc.save(output_path)


def main():
    if len(sys.argv) != 3:
        print("Usage: python scripts/generate_designed_cv.py input.json output.docx", file=sys.stderr)
        sys.exit(2)
    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    data = json.loads(input_path.read_text(encoding="utf-8"))
    build_doc(data, output_path)
    print(output_path)


if __name__ == "__main__":
    main()
