"""
BoD formatting for "Think – Ask – Click"
Format: 12x19cm, Palatino Linotype, 9pt body text
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document('/home/ubuntu/upload/BoD-Microsoft-Word-Vorlage-12x19cm.docx')

# Clear all existing paragraphs
for para in doc.paragraphs:
    p = para._element
    p.getparent().remove(p)

# Page setup
section = doc.sections[0]
section.page_width = Cm(12)
section.page_height = Cm(19)
section.top_margin = Cm(1.2)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(1.7)
section.right_margin = Cm(1.5)
section.header_distance = Cm(1.25)
section.footer_distance = Cm(1.25)

FONT_NAME = 'Palatino Linotype'
BODY_SIZE = Pt(9)
H1_SIZE = Pt(13)
H2_SIZE = Pt(11)
H3_SIZE = Pt(9.5)

def set_font(run, name=FONT_NAME, size=BODY_SIZE, bold=False, italic=False):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic

def add_paragraph(doc, space_before=Pt(0), space_after=Pt(4),
                  line_spacing=Pt(12), align=WD_ALIGN_PARAGRAPH.LEFT,
                  first_indent=None, left_indent=None):
    para = doc.add_paragraph()
    para.style = doc.styles['Normal']
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    return para

def add_run(para, text, bold=False, italic=False, size=BODY_SIZE):
    run = para.add_run(text)
    set_font(run, bold=bold, italic=italic, size=size)
    return run

def insert_page_break(doc):
    para = doc.add_paragraph()
    para.style = doc.styles['Normal']
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def parse_inline(para, text, base_size=BODY_SIZE):
    pattern = re.compile(r'(\*\*.*?\*\*|\*[^*]+\*)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            add_run(para, part[2:-2], bold=True, size=base_size)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            add_run(para, part[1:-1], italic=True, size=base_size)
        else:
            add_run(para, part, size=base_size)

# ===== FRONT MATTER =====

# Page 1: Half-title
p = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Cm(3), space_after=Pt(6))
add_run(p, 'Roger Basler de Roca', bold=True, size=Pt(9))

p = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(4))
add_run(p, 'Think – Ask – Click', bold=True, size=Pt(11))

insert_page_break(doc)

# Page 2: Blank
p = add_paragraph(doc, space_before=Pt(0), space_after=Pt(0))
add_run(p, '')
insert_page_break(doc)

# Page 3: Full title page
p = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Cm(2), space_after=Pt(6))
add_run(p, 'Roger Basler de Roca', size=Pt(11))

p = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(18), space_after=Pt(8))
add_run(p, 'Think – Ask – Click', bold=True, size=Pt(18))

p = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(4))
add_run(p, "Who's Really Thinking When ChatGPT Answers?", italic=True, size=Pt(10))

p = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(4), space_after=Pt(4))
add_run(p, 'A Practical Guide to Independent Thinking in the Digital Age', size=Pt(9))

insert_page_break(doc)

# Page 4: Copyright / Imprint
p = add_paragraph(doc, space_before=Pt(0), space_after=Pt(6))
add_run(p, 'Copyright', bold=True, size=Pt(9))

imprint_lines = [
    'Bibliographic information of the Deutsche Nationalbibliothek:',
    'The Deutsche Nationalbibliothek lists this publication in the Deutsche',
    'Nationalbibliografie; detailed bibliographic data are available on the',
    'Internet at http://dnb.d-nb.de.',
    '',
    '© 2026 Roger Basler de Roca',
    'All rights reserved.',
    '',
    'Publisher: BoD – Books on Demand GmbH, Überseering 33, 22297 Hamburg, Germany',
    'Print: Libri Plureos GmbH, Friedensallee 273, 22763 Hamburg, Germany',
    '',
    'ISBN: 978-3-XXXX-XXXX-X',
    '',
    'No part of this publication may be reproduced, stored in a retrieval system,',
    'or transmitted in any form or by any means without the prior written',
    'permission of the author.',
]
for line in imprint_lines:
    p = add_paragraph(doc, space_before=Pt(0), space_after=Pt(2))
    add_run(p, line, size=Pt(8))

insert_page_break(doc)

# ===== BOOK CONTENT =====

with open('/home/ubuntu/think_ask_click_english.md', 'r') as f:
    md_text = f.read()

lines = md_text.split('\n')

in_list = False
list_buffer = []

def flush_list(doc, items):
    for item in items:
        p = add_paragraph(doc, space_before=Pt(0), space_after=Pt(3),
                          line_spacing=Pt(12), left_indent=Cm(0.4), first_indent=Cm(-0.4))
        add_run(p, '– ', size=BODY_SIZE)
        parse_inline(p, item.strip(), base_size=BODY_SIZE)

i = 0
while i < len(lines):
    line = lines[i]
    raw = line.rstrip()

    if raw.strip() == '':
        if in_list:
            flush_list(doc, list_buffer)
            list_buffer = []
            in_list = False
        i += 1
        continue

    if raw.strip().startswith('---'):
        i += 1
        continue

    # H1: ## (chapters)
    if raw.startswith('## ') and not raw.startswith('### '):
        if in_list:
            flush_list(doc, list_buffer); list_buffer = []; in_list = False
        insert_page_break(doc)
        p = add_paragraph(doc, space_before=Pt(54), space_after=Pt(18),
                          line_spacing=Pt(16))
        add_run(p, raw[3:], bold=True, size=H1_SIZE)
        i += 1
        continue

    # H2: ###
    if raw.startswith('### ') and not raw.startswith('#### '):
        if in_list:
            flush_list(doc, list_buffer); list_buffer = []; in_list = False
        p = add_paragraph(doc, space_before=Pt(14), space_after=Pt(4),
                          line_spacing=Pt(13))
        add_run(p, raw[4:], bold=True, size=H2_SIZE)
        i += 1
        continue

    # H3: ####
    if raw.startswith('#### '):
        if in_list:
            flush_list(doc, list_buffer); list_buffer = []; in_list = False
        p = add_paragraph(doc, space_before=Pt(10), space_after=Pt(3),
                          line_spacing=Pt(12))
        add_run(p, raw[5:], bold=True, italic=True, size=H3_SIZE)
        i += 1
        continue

    # Top-level #
    if raw.startswith('# ') and not raw.startswith('## '):
        if in_list:
            flush_list(doc, list_buffer); list_buffer = []; in_list = False
        insert_page_break(doc)
        p = add_paragraph(doc, space_before=Pt(54), space_after=Pt(18),
                          line_spacing=Pt(16))
        add_run(p, raw[2:], bold=True, size=H1_SIZE)
        i += 1
        continue

    # List items
    if raw.lstrip().startswith('- ') or raw.lstrip().startswith('* '):
        in_list = True
        item_text = re.sub(r'^[\s\-\*]+', '', raw)
        list_buffer.append(item_text)
        i += 1
        continue

    if in_list:
        flush_list(doc, list_buffer); list_buffer = []; in_list = False

    # Body paragraph
    p = add_paragraph(doc, space_before=Pt(0), space_after=Pt(4),
                      line_spacing=Pt(12), first_indent=Cm(0.4))
    parse_inline(p, raw, base_size=BODY_SIZE)
    i += 1

if in_list:
    flush_list(doc, list_buffer)

# ===== REFERENCES SECTION =====
insert_page_break(doc)
p = add_paragraph(doc, space_before=Pt(54), space_after=Pt(18), line_spacing=Pt(16))
add_run(p, 'References', bold=True, size=H1_SIZE)

references = [
    "Chirayath, G. (2025). Cognitive offloading or cognitive overload? How AI alters the mind. PMC. https://pmc.ncbi.nlm.nih.gov/articles/PMC12678390/",
    "Duke Learning Innovation. (2024). Does AI harm critical thinking? https://lile.duke.edu/ai-ethics-learning-toolkit/does-ai-harm-critical-thinking/",
    "Frontiers in Education. (2025). Evaluating the impact of AI on critical thinking skills. https://www.frontiersin.org/journals/education/articles/10.3389/feduc.2025.1719625/full",
    "Gerlich, M. (2025). AI tools in society: Impacts on cognitive offloading and the future of critical thinking. Societies, 15(1), 6. https://doi.org/10.3390/soc15010006",
    "Kosmyna, N., et al. (2025). Your brain on ChatGPT: Accumulation of cognitive debt when using an AI assistant for essay writing task. arXiv preprint arXiv:2506.08872. https://arxiv.org/abs/2506.08872",
    "MIT Media Lab / Time Magazine. (2025, June 23). ChatGPT's impact on our brains according to an MIT study. Time. https://time.com/7295195/ai-chatgpt-google-learning-school/",
    "Sparrow, B., Liu, J., & Wegner, D. M. (2011). Google effects on memory: Cognitive consequences of having information at our fingertips. Science, 333(6043), 776–778. https://doi.org/10.1126/science.1207745",
    "Stanford SCALE Lab. (2025). ChatGPT produces more 'lazy' thinkers: Evidence of cognitive engagement decline. https://scale.stanford.edu/ai/repository/chatgpt-produces-more-lazy-thinkers-evidence-cognitive-engagement-decline",
]

for ref in references:
    p = add_paragraph(doc, space_before=Pt(0), space_after=Pt(4),
                      line_spacing=Pt(12), first_indent=Cm(-0.5), left_indent=Cm(0.5))
    add_run(p, ref, size=Pt(8))

# ===== FOOTER with page numbers =====
for section in doc.sections:
    footer = section.footer
    for para in footer.paragraphs:
        para.clear()
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    set_font(run, size=Pt(8))
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

doc.save('/home/ubuntu/Think_Ask_Click_BoD.docx')
print("English BoD document created successfully!")
