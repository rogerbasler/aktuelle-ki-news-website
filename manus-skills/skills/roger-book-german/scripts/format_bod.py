"""
BoD-Formatierung für "Denk - Frag - Klick"
Format: 12x19cm
Schrift: Palatino Linotype, 9pt Fliesstext
Ränder: oben 1.2cm, unten 2.0cm, links 1.7cm, rechts 1.5cm
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import copy

# Vorlage laden
template = Document('/home/ubuntu/upload/BoD-Microsoft-Word-Vorlage-12x19cm.docx')

# Neues Dokument auf Basis der Vorlage
doc = Document('/home/ubuntu/upload/BoD-Microsoft-Word-Vorlage-12x19cm.docx')

# Alle bestehenden Absätze löschen
for para in doc.paragraphs:
    p = para._element
    p.getparent().remove(p)

# Seitenformat sicherstellen
section = doc.sections[0]
section.page_width = Cm(12)
section.page_height = Cm(19)
section.top_margin = Cm(1.2)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(1.7)
section.right_margin = Cm(1.5)
section.header_distance = Cm(1.25)
section.footer_distance = Cm(1.25)

FONT_NAME = 'Palatino Linotype'
BODY_SIZE = Pt(9)
H1_SIZE = Pt(13)
H2_SIZE = Pt(11)
H3_SIZE = Pt(10)

def set_font(run, name=FONT_NAME, size=BODY_SIZE, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_paragraph(doc, text='', style='Normal', align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=Pt(0), space_after=Pt(4),
                  line_spacing=Pt(12), first_indent=None, left_indent=None):
    para = doc.add_paragraph()
    para.style = doc.styles['Normal']
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    return para

def add_run_with_font(para, text, bold=False, italic=False, size=BODY_SIZE):
    run = para.add_run(text)
    set_font(run, bold=bold, italic=italic, size=size)
    return run

def add_page_break(doc):
    para = doc.add_paragraph()
    para.style = doc.styles['Normal']
    run = para.add_run()
    run.add_break(docx.oxml.ns.qn('w:br'))
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)

# Seitenumbruch via XML
def insert_page_break(doc):
    para = doc.add_paragraph()
    para.style = doc.styles['Normal']
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

# Markdown-Text in Runs aufteilen (Bold, Italic)
def parse_inline(para, text, base_size=BODY_SIZE):
    # Verarbeite **bold**, *italic*, und normalen Text
    pattern = re.compile(r'(\*\*.*?\*\*|\*[^*]+\*)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            add_run_with_font(para, part[2:-2], bold=True, size=base_size)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            add_run_with_font(para, part[1:-1], italic=True, size=base_size)
        else:
            add_run_with_font(para, part, size=base_size)

# Buchtext laden
with open('/home/ubuntu/denk_frag_klick_final_v2.md', 'r') as f:
    md_text = f.read()

lines = md_text.split('\n')

# ===== TITELEI =====

# Seite 1: Schmutztitel
p = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Cm(3), space_after=Pt(6))
add_run_with_font(p, 'Roger Basler de Roca', bold=True, size=Pt(9))

p = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(4))
add_run_with_font(p, 'Denk – Frag – Klick', bold=True, size=Pt(11))

insert_page_break(doc)

# Seite 2: Leerseite
p = add_paragraph(doc, space_before=Pt(0), space_after=Pt(0))
add_run_with_font(p, '')

insert_page_break(doc)

# Seite 3: Haupttitelseite
p = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Cm(2), space_after=Pt(6))
add_run_with_font(p, 'Roger Basler de Roca', size=Pt(11))

p = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(18), space_after=Pt(8))
add_run_with_font(p, 'Denk – Frag – Klick', bold=True, size=Pt(18))

p = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(4))
add_run_with_font(p, 'Wer denkt, wenn ChatGPT antwortet?', italic=True, size=Pt(10))

p = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(4), space_after=Pt(4))
add_run_with_font(p, 'Ein Praxisguide für selbstständiges Denken im digitalen Alltag', size=Pt(9))

insert_page_break(doc)

# Seite 4: Impressum
p = add_paragraph(doc, space_before=Pt(0), space_after=Pt(6))
add_run_with_font(p, 'Impressum', bold=True, size=Pt(9))

impressum_lines = [
    'Bibliografische Information der Deutschen Nationalbibliothek:',
    'Die Deutsche Nationalbibliothek verzeichnet diese Publikation in der',
    'Deutschen Nationalbibliografie; detaillierte bibliografische Daten sind',
    'im Internet über http://dnb.d-nb.de abrufbar.',
    '',
    '© 2026 Roger Basler de Roca',
    '',
    'Verlag: BoD – Books on Demand GmbH, Überseering 33, 22297 Hamburg',
    'Druck: Libri Plureos GmbH, Friedensallee 273, 22763 Hamburg',
    '',
    'ISBN: 978-3-XXXX-XXXX-X',
    '',
    'Alle Rechte vorbehalten. Kein Teil dieses Werkes darf ohne schriftliche',
    'Genehmigung des Autors reproduziert, vervielfältigt oder verbreitet werden.',
]
for line in impressum_lines:
    p = add_paragraph(doc, space_before=Pt(0), space_after=Pt(2))
    add_run_with_font(p, line, size=Pt(8))

insert_page_break(doc)

# ===== BUCHINHALT =====

in_list = False
list_buffer = []

def flush_list(doc, items):
    for item in items:
        p = add_paragraph(doc, space_before=Pt(0), space_after=Pt(3),
                          line_spacing=Pt(12), left_indent=Cm(0.4), first_indent=Cm(-0.4))
        add_run_with_font(p, '– ', bold=False, size=BODY_SIZE)
        parse_inline(p, item.strip(), base_size=BODY_SIZE)

i = 0
while i < len(lines):
    line = lines[i]
    raw = line.rstrip()

    # Leere Zeile
    if raw.strip() == '':
        if in_list:
            flush_list(doc, list_buffer)
            list_buffer = []
            in_list = False
        i += 1
        continue

    # Seitenumbruch-Marker
    if raw.strip().startswith('---'):
        i += 1
        continue

    # Heading 1 (##)
    if raw.startswith('## ') and not raw.startswith('### '):
        if in_list:
            flush_list(doc, list_buffer)
            list_buffer = []
            in_list = False
        insert_page_break(doc)
        p = add_paragraph(doc, space_before=Pt(54), space_after=Pt(18),
                          line_spacing=Pt(16), align=WD_ALIGN_PARAGRAPH.LEFT)
        add_run_with_font(p, raw[3:], bold=True, size=H1_SIZE)
        i += 1
        continue

    # Heading 2 (###)
    if raw.startswith('### ') and not raw.startswith('#### '):
        if in_list:
            flush_list(doc, list_buffer)
            list_buffer = []
            in_list = False
        p = add_paragraph(doc, space_before=Pt(14), space_after=Pt(4),
                          line_spacing=Pt(13), align=WD_ALIGN_PARAGRAPH.LEFT)
        add_run_with_font(p, raw[4:], bold=True, size=H2_SIZE)
        i += 1
        continue

    # Heading 3 (####)
    if raw.startswith('#### '):
        if in_list:
            flush_list(doc, list_buffer)
            list_buffer = []
            in_list = False
        p = add_paragraph(doc, space_before=Pt(10), space_after=Pt(3),
                          line_spacing=Pt(12), align=WD_ALIGN_PARAGRAPH.LEFT)
        add_run_with_font(p, raw[5:], bold=True, italic=True, size=Pt(9.5))
        i += 1
        continue

    # Heading # (Vorwort-Ebene, wird wie H1 behandelt)
    if raw.startswith('# ') and not raw.startswith('## '):
        if in_list:
            flush_list(doc, list_buffer)
            list_buffer = []
            in_list = False
        insert_page_break(doc)
        p = add_paragraph(doc, space_before=Pt(54), space_after=Pt(18),
                          line_spacing=Pt(16), align=WD_ALIGN_PARAGRAPH.LEFT)
        add_run_with_font(p, raw[2:], bold=True, size=H1_SIZE)
        i += 1
        continue

    # Aufzählung (- oder *)
    if raw.lstrip().startswith('- ') or raw.lstrip().startswith('* '):
        in_list = True
        item_text = re.sub(r'^[\s\-\*]+', '', raw)
        list_buffer.append(item_text)
        i += 1
        continue

    # Normaler Absatz
    if in_list:
        flush_list(doc, list_buffer)
        list_buffer = []
        in_list = False

    p = add_paragraph(doc, space_before=Pt(0), space_after=Pt(4),
                      line_spacing=Pt(12), first_indent=Cm(0.4))
    parse_inline(p, raw, base_size=BODY_SIZE)
    i += 1

# Letzte Liste leeren
if in_list:
    flush_list(doc, list_buffer)

# Fusszeile mit Seitenzahl
for section in doc.sections:
    footer = section.footer
    for para in footer.paragraphs:
        para.clear()
    if footer.paragraphs:
        fp = footer.paragraphs[0]
    else:
        fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    set_font(run, size=Pt(8))
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

doc.save('/home/ubuntu/Denk_Frag_Klick_BoD.docx')
print("BoD-Dokument erfolgreich erstellt!")
print(f"Seiten (geschätzt): {len([p for p in doc.paragraphs if p.text.strip()])}")
